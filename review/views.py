# -*- coding: utf-8 -*-
from __future__ import unicode_literals
from __future__ import print_function

from django.core.mail import EmailMessage
from django.shortcuts import render
from django.template.loader import get_template, render_to_string
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt

from docx import Document
from io import BytesIO
from xhtml2pdf import pisa


from .forms import ReviewForm, UserInfoForm
from .models import ReviewDraft

import json
import uuid
import traceback
import re


# Function to generate a DOCX document from the context dictionary
def generate_docx(context_dict):
    doc = Document()
    doc.add_heading('PDS Data Set Peer Review', 0)

   # Helper function to add a bold label and normal text value
    def add_bold_label(label, value):
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        p.add_run(value)

    # Helper function to add a bold question and normal answer
    def add_bold_question_with_answer(question, answer):
        p = doc.add_paragraph()
        run_q = p.add_run(question + '\n')
        run_q.bold = True
        p.add_run(answer)

    # Add contact info
    add_bold_label("Name", context_dict.get('contact_name', ''))
    add_bold_label("Email", context_dict.get('contact_email', ''))
    add_bold_label("Reviewed PDS Data Set", context_dict.get('derived_data', ''))

    # Add questions with responses
    add_bold_question_with_answer(
        "Does the data set/bundle provide clear and concise documentation adequate for its usage?",
        context_dict.get('question1', '')
    )
    add_bold_question_with_answer(
        "Are you able to manipulate and/or plot the data, interpret columns/rows into tables, and understand the context and relationships of the data products?",
        context_dict.get('question2', '')
    )
    add_bold_question_with_answer(
        "Are there any concerns about the creation/generation, calibration, or general usability of the data?",
        context_dict.get('question3', '')
    )
    add_bold_question_with_answer(
        "Were there any issues with the data access website, related references, or any other accessibility concerns?",
        context_dict.get('question4', '')
    )
    add_bold_question_with_answer(
        "Do you have any further comments to the PDS Atmospheres Node about the data?",
        context_dict.get('question5', '')
    )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Function to generate a PDF document from the context dictionary
def generate_pdf(context_dict):
    html = render_to_string('review/comment_template.txt', context_dict)
    buffer = BytesIO()
    pisa.CreatePDF(src=html, dest=buffer)
    buffer.seek(0)
    return buffer


# Index view: It handles both GET and POST requests, processes the review form, and sends emails
def index(request):
    request.encoding = 'utf-8'

    draft_id = request.COOKIES.get('draft_id')
    initial_data = {}

    # Try to load draft data if draft_id cookie exists
    if draft_id:
        try:
            draft_uuid = uuid.UUID(draft_id)
            draft = ReviewDraft.objects.get(draft_id=draft_uuid)
            if draft.content:
                initial_data = json.loads(draft.content)
        except (ValueError, ReviewDraft.DoesNotExist, json.JSONDecodeError):
            initial_data = {}

    # Pass initial data as 'initial' to pre-fill form
    review_form = ReviewForm(request.POST or None, initial=initial_data)
    user_info_form = UserInfoForm(request.POST or None)

    context_dict = {
        'review_form': review_form,
        'user_info_form': user_info_form,
        'email_sent': False,
        'draft_id': draft_id or '',
    }

    if review_form.is_valid():
        context_dict['contact_name'] = '{0}'.format(review_form.cleaned_data['user_name'])
        context_dict['contact_email'] = review_form.cleaned_data['user_email']
        context_dict['derived_data'] = review_form.cleaned_data['derived_data']
        context_dict['question1'] = review_form.cleaned_data['question1']
        context_dict['question2'] = review_form.cleaned_data['question2']
        context_dict['question3'] = review_form.cleaned_data['question3']
        context_dict['question4'] = review_form.cleaned_data['question4']
        context_dict['question5'] = review_form.cleaned_data['question5']
        context_dict['recipient'] = review_form.cleaned_data['recipient']

        # Find template used for email confirmation
        template = get_template('review/comment_template.txt')
        content = template.render(context_dict)

        # Generate attachments
        docx_file = generate_docx(context_dict)
        pdf_file = generate_pdf(context_dict)

        email = EmailMessage(
            subject="PDS Data Set Peer Review from {}".format(context_dict['contact_name']),
            body=(
                "A new review has been submitted by {name}.\n\n"
                "Reviewed Data Set: {data}\n"
                "Recipient: {recipient}\n\n"
                "Please find the attached documents for details."
            ).format(
                name=context_dict['contact_name'],
                data=context_dict['derived_data'],
                recipient="Lynn Neakrase" if context_dict['recipient'] == "lynn" else "Lyle Huber"
            ),
            from_email='atm-elsa@nmsu.edu',
            to=['rupakdey@nmsu.edu', 'lneakras@nmsu.edu', 'sajomont@nmsu.edu', 'lhuber@nmsu.edu'],
            #to=['rupakdey@nmsu.edu'],
            headers={'Reply-To': context_dict['contact_email']}
        )

        #Naming the files based on the Reviewed Data Set
        def sanitize_filename(name):
            # Lowercase, replace spaces with underscores, remove non-alphanumeric/underscore chars
            safe_name = name.lower()
            safe_name = safe_name.replace(' ', '_')
            safe_name = re.sub(r'[^a-z0-9_]+', '', safe_name)
            return safe_name or "review"
        
        derived_data_name = sanitize_filename(context_dict.get('derived_data', 'review'))

        docx_filename = f"{derived_data_name}_review.docx"
        pdf_filename = f"{derived_data_name}_review.pdf"

        # Attach the generated DOCX and PDF files
        email.attach(docx_filename, docx_file.read(), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        email.attach(pdf_filename, pdf_file.read(), 'application/pdf')
        print("Before sending email to ELSA team")
        email.send()
        print("After sending email to ELSA team")

        email_confirmation = EmailMessage(
            subject="Thank you for submitting a PDS Data Set Peer Review!",
            body="Your review has been received. A copy of the review is included here for your record. Please find the attachments! \nThank you for using ELSA!\n\nRegards,\nTeam ELSA",
            from_email='atm-elsa@nmsu.edu',
            to=[context_dict['contact_email']]
        )

        # Attach the same DOCX and PDF files to the confirmation email
        email_confirmation.attach(docx_filename, docx_file.getvalue(), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        email_confirmation.attach(pdf_filename, pdf_file.getvalue(), 'application/pdf')
        print("Before sending email confirmation to user")
        email_confirmation.send()
        print("After sending email")

        # Clear the draft after successful submission
        if draft_id:
            try:
                draft_uuid = uuid.UUID(draft_id)
                ReviewDraft.objects.filter(draft_id=draft_uuid).delete()
            except ValueError:
                pass  # Invalid UUID, silently skip
            
        context_dict['email_sent'] = True
        return render(request, 'review/index.html', context_dict)

    return render(request, 'review/index.html', context_dict)

# Save draft view: It accepts a POST request with JSON data, saves it as a draft, and returns a JSON response
@csrf_exempt
def save_draft(request, draft_id):
    if request.method != 'POST':
        return JsonResponse({'error': 'Only POST allowed'}, status=405)

    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)

    # Validate UUID format
    try:
        draft_uuid = uuid.UUID(draft_id)
    except ValueError:
        return JsonResponse({'error': 'Invalid draft ID'}, status=400)

    try:
        draft, created = ReviewDraft.objects.get_or_create(draft_id=draft_uuid)

        # Ensure content is a JSON string before saving
        if isinstance(data, str):
            # if string, validate it's valid JSON string
            try:
                json.loads(data)
                draft.content = data
            except json.JSONDecodeError:
                return JsonResponse({'error': 'Invalid JSON string'}, status=400)
        else:
            draft.content = json.dumps(data)

        if request.user.is_authenticated and not draft.user:
            draft.user = request.user

        draft.save()
    except Exception:
        traceback.print_exc()
        return JsonResponse({'error': 'Internal server error'}, status=500)

    return JsonResponse({'status': 'saved', 'created': created})


# Load draft view: It retrieves the draft by ID and returns its content as JSON
def load_draft(request, draft_id):
    try:
        draft_uuid = uuid.UUID(draft_id)
    except ValueError:
        return JsonResponse({'error': 'Invalid draft ID'}, status=400)

    try:
        draft = ReviewDraft.objects.get(draft_id=draft_uuid)
        content = json.loads(draft.content) if draft.content else {}
        return JsonResponse(content)
    except ReviewDraft.DoesNotExist:
        return JsonResponse({}, status=404)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Stored draft content corrupted'}, status=500)
